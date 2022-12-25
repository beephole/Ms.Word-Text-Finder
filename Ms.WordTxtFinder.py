import argparse, re, os, docx, time, sys, fnmatch, getpass
from docx import Document
import tkinter as tk
from tkinter import filedialog
import docx2pdf


template = """
___  ___         _    _                   _   _____        _____  ______  _             _             
|  \/  |        | |  | |                 | | |_   _|      |_   _| |  ___|(_)           | |            
| .  . | ___    | |  | |  ___   _ __   __| |   | |  __  __  | |   | |_    _  _ __    __| |  ___  _ __ 
| |\/| |/ __|   | |/\| | / _ \ | '__| / _` |   | |  \ \/ /  | |   |  _|  | || '_ \  / _` | / _ \| '__|
| |  | |\__ \ _ \  /\  /| (_) || |   | (_| |   | |   >  <   | |   | |    | || | | || (_| ||  __/| |   
\_|  |_/|___/(_) \/  \/  \___/ |_|    \__,_|   \_/  /_/\_\  \_/   \_|    |_||_| |_| \__,_| \___||_|    
Copyright (c) 2022 Beephole
This software is licensed under the MIT License. See the LICENSE file for details.                               
"""
print(template, end="\n")

parser = argparse.ArgumentParser(
    description="Ms.Word Text Finder is a Python program that allows users to search for and extract specific text from a Microsoft Word document (.docx). Once the file is open, you can enter a search term to search for within the document. If the term is found, the program will print the paragraph containing the term and allow you to specify points within the paragraph to extract a portion of the text. The extracted text can be saved as a variable and written to a text file."
)
parser.add_argument(
    "-f",
    "--file",
    required=False,
    help="File name to be used. Ex 'wordDocument' or 'wordDocument.docx ",
)
onedrive_path = os.path.join(os.path.expanduser("~"), "OneDrive")
if os.path.exists(onedrive_path):
    default_directory = onedrive_path
else:
    default_directory = os.path.join(os.path.expanduser("~"), "Desktop")
parser.add_argument(
    "-d",
    "--directory",
    default=default_directory,
    required=False,
    help="directory to search for the file (default: desktop)",
)
parser.add_argument(
    "-s",
    "--scann",
    nargs="?",
    const=True,
    default=False,
    help="Scan the template for variables",
)
parser.add_argument(
    "-t",
    "--template",
    required=False,
    help="template file name to be used. Ex 'template' or 'template.docx'",
)
parser.add_argument(
    "-o",
    "--output",
    required=False,
    help="output file name. Ex 'output' or 'output.txt'",
)
parser.add_argument(
    "-i",
    "--input",
    required=False,
    help="input file name to be used Ex 'inputfile' or 'inputfile.txt'",
)
parser.add_argument(
    "-H",
    "--help-message",
    action="help",
    help="Note: When setting a word as a point and the word may include a symbol without a space, such as in the string 'last name:', be sure to include the symbol in your selection. For example, to obtain the index of the word 'name', you should include the symbol ':'  If there is a space between the word and the symbol, it is acceptable to simply select the word. Alternatively, you can also split the string by the symbol. Please exercise caution when making these selectionsAND CLOSE THE CURRENTS WORD DOCUMENT WHEN WORKING WITH IT ",
)
parser.add_argument(
    "-b",
    "--browse",
    action="store_true",
    help="Open a Tk window to browse for a file",
)
parser.add_argument("--bold", action="store_true", help="make the replaced text bold")

parser.add_argument(
    "-pdf", "--pdf", action="store_true", help="Convert the template to a PDF file"
)


args = parser.parse_args()


filename = args.file
search_directory = args.directory
template_filename = args.template
output_filename = args.output
browser_tk = args.browse
conver_pdf = args.pdf
input_var = args.input
bold = args.bold
i = 0

results = {}


def search_text():
    while True:
        print("\n")
        new_search = input("Enter the Text that You want to search in the Document: ")
        if new_search:
            print("\n")
            return new_search
        else:
            print("Please enter a valid search term.")
            print("\n")


def loop_text(document, user_input):
    for i, paragraph in enumerate(document.paragraphs):
        if user_input in paragraph.text:
            paragraph_index10 = i
            line = document.paragraphs[paragraph_index10].text.split()
            print("Here it is What you Looked for: ")
            print("\n")
            print(" ".join(line))

            return line


def searchMsWord():
    newTextList = []
    results = {}
    while True:
        answ = input("Do you want to search Text from a Ms.Word doc? (yes/no):")
        print("\n")
        if answ.lower() not in ["yes", "no"] or answ.strip() == "":
            print("Invalid input. Please type something'")
            print("\n")
            continue

        elif answ.lower() == "yes":
            user_input = search_text()

            if user_input.strip() != "":
                loop_text(document, user_input)
                print("\n")
                while True:
                    answ = input("Are you happy with the Result? (yes/no/q): ")
                    print("\n")
                    if answ.lower() not in ["yes", "no", "q"]:
                        print("Invalid input. Please enter 'yes', 'no', or 'q'.")
                        print("\n")
                        continue
                    if answ.lower() == "yes":
                        line = loop_text(document, user_input)
                        print("\n")
                        print("Ok thats Great !!")
                        print("\n")

                        time.sleep(1)
                        while True:
                            try:
                                question = int(
                                    input(
                                        "Do you want to cut the string at 1 Point or 2 Points (1/2): "
                                    )
                                )
                                print("\n")
                                print(
                                    "Remember if the word you want to grab ends with a symbol , grab the symbol also !"
                                )
                                print("\n")
                                if question not in [1, 2]:
                                    print(
                                        "Invalid input. Please enter either 1 or 2. (1/2)"
                                    )
                                    print("\n")
                                    continue
                                break
                            except ValueError:
                                print(
                                    "Invalid input. Please enter either 1 or 2.  (1/2)"
                                )
                                print("\n")
                        if question == 1:

                            print(
                                "NOTE -> For cutting left or right all we need is 1 Point",
                                end="\n",
                            )
                            print("\n")
                            whereToCut = ""
                            while not whereToCut:
                                whereToCut = input(
                                    "Enter the  Point you want to cut it -> Should be symbol or word :  "
                                )
                                print("\n")
                                if not whereToCut:
                                    print(
                                        "Please enter a valid point. The point cannot be an empty string."
                                    )
                                    print("\n")

                            whichSideCutt = ""
                            while whichSideCutt.lower() not in [
                                "left",
                                "right",
                                "double",
                            ]:
                                whichSideCutt = input(
                                    "Choose you want to cut and keep the 'left' or 'right' of the point? (left/right): "
                                )
                                print("\n")
                                if whichSideCutt.lower() not in [
                                    "left",
                                    "right",
                                    "double",
                                ]:
                                    print(
                                        "Please enter a valid direction (left/right/double)."
                                    )
                                    print("\n")
                                elif whichSideCutt == "right":
                                    try:
                                        indexOfwhereToCut = line.index(whereToCut)
                                        line = line[indexOfwhereToCut + 1 :]
                                    except ValueError:
                                        print(
                                            f"The point '{whereToCut}' was not found in the string '{line}'."
                                        )
                                        break
                                    newTextList.append(line)
                                elif whichSideCutt == "left":
                                    try:
                                        indexOfwhereToCut = line.index(whereToCut)
                                        line = line[:indexOfwhereToCut]
                                        newTextList.append(line)
                                    except ValueError:
                                        print(
                                            f"The point '{whereToCut}' was not found in the string '{line}'."
                                        )
                                        break
                        elif question == 2:
                            whereToCut1 = ""
                            while not whereToCut1:
                                whereToCut1 = input(
                                    "Enter the  1st Point you want to cut it ->Should be symbol or word : "
                                )
                                if not whereToCut1:
                                    print(
                                        "Please enter a valid point. The point cannot be an empty string."
                                    )
                            whereToCut2 = ""
                            while not whereToCut2:
                                whereToCut2 = input(
                                    "Enter the 2nd Point you want to cut it ->Should be symbol or word : "
                                )
                                if not whereToCut2:
                                    print(
                                        "Please enter a valid point. The point cannot be an empty string."
                                    )
                            print("\n")
                            try:
                                indexOfwhereToCut1 = line.index(whereToCut1)
                                indexOfwhereToCut2 = line.index(whereToCut2)
                            except ValueError:
                                print(
                                    f"The point '{whereToCut1}'and/or '{whereToCut2}' was not found in the string '{line}'."
                                )
                                continue
                            if indexOfwhereToCut2 > indexOfwhereToCut1:
                                line = line[indexOfwhereToCut1 + 1 : indexOfwhereToCut2]
                                newTextList.append(line)
                            else:
                                indexOfwhereToCut1, indexOfwhereToCut2 = (
                                    indexOfwhereToCut2,
                                    indexOfwhereToCut1,
                                )
                                line = line[indexOfwhereToCut1 + 1 : indexOfwhereToCut2]
                                newTextList.append(line)

                        else:
                            print("Something went wrong,Try to type a number : 1 or 2")
                            continue

                        ######
                        print("Here it is the result ---------> " + " ".join(line))
                        print("\n")
                        if len(newTextList) > 0:
                            RRezultatit = newTextList[0]
                            return RRezultatit
                        else:
                            print(
                                f"The point '{whereToCut}' was not found in the string '{line}'."
                            )
                    elif answ.lower() == "no":
                        print("Sorry, Try Again .....")
                        time.sleep(1)
                        results.pop(variable_name, None)

                        break

                    elif answ.lower() == "q":
                        print("Exiting .....")
                        print("\n")
                        print("Have a Good Day  !!")
                        sys.exit()

            elif user_input.strip() == "":
                print("Invalid input. Please type something")
                results.pop(variable_name, None)
                time.sleep(1)
                break
            continue
        elif answ.lower() == "no":
            prgj = ""
            while not prgj:
                prgj = input("Enter a value for the variable: ")
                if not prgj:
                    print("Enter something fam!")
            value = prgj
            newTextList.append(value)
            return newTextList
        else:
            break
        break


def is_valid_variable_name(name):
    """Returns True if the given string is a valid variable name, False otherwise."""
    if not re.match(r"^\w+$", name):
        return False

    if name[0].isdigit():
        return False
    if name in [
        "and",
        "as",
        "assert",
        "async",
        "await",
        "break",
        "class",
        "continue",
        "def",
        "del",
        "elif",
        "else",
        "except",
        "False",
        "finally",
        "for",
        "from",
        "global",
        "if",
        "import",
        "in",
        "is",
        "lambda",
        "None",
        "nonlocal",
        "not",
        "or",
        "pass",
        "raise",
        "return",
        "True",
        "try",
        "while",
        "with",
        "yield",
    ]:
        return False

    return True


if args.browse:
    
    root = tk.Tk()
    root.withdraw() 
    browsefile_path = filedialog.askopenfilename()
    browsefilename = browsefile_path.split("/")[
        -1
    ] 

    if not browsefilename.endswith(".docx"):
        browsefilename += ".docx"
    for root, dirs, files in os.walk(search_directory):
        if browsefilename in files:
            browseile_path = os.path.join(root, browsefilename)
            browsefile_path = os.path.join(root, browsefilename)
            browsefile_variable = browsefile_path
            break
    else:
        print(
            f"Error: {filename} was not found in directory on the specified{search_directory}"
        )
        print("Check the filename or check the path if it's written correctly ! ")
        print("Exiting ...")
        sys.exit()
    document_name = browsefile_variable
    try:
        print(document_name, end="\n")
        document = docx.Document(document_name)
        print("\n")
    except Exception:
        print("Error: Invalid file path")


if args.file:
    if not filename.endswith(".docx"):
        filename += ".docx"

    for root, dirs, files in os.walk(search_directory):
        if filename in files:
            file_path = os.path.join(root, filename)
            file_variable = file_path
            break
    else:
        print(
            f"Error: {filename} was not found in directory on the specified{search_directory}"
        )
        print("Check the filename or check the path if it's written correctly ! ")
        print("Exiting ...")
        sys.exit()
    document_name = file_variable
    try:
        print(document_name, end="\n")
        document = docx.Document(document_name)
        print("\n")
    except Exception:
        print("Error: Invalid file path")


if args.input:
    results1 = {}
    if not input_var.endswith(".txt"):
        input_var += ".txt"
    for root, dirs, files in os.walk(search_directory):
        if input_var in files:
            input_file_path = os.path.join(root, input_var)
            input_file_variable = input_file_path
            break
    else:
        print(
            f"Error: {input_var} was not found in directory on the specified {search_directory}"
        )
        print("Check the filename or check the path if it's written correctly ! ")
        print("Exiting ...")
        sys.exit()

    input_document = input_file_variable
    try:
        print(input_document, end="\n")

    except Exception:
        print("Error: Invalid file path")

    with open(input_document, "r") as f:
        contents = f.read()

    for line in contents.split("\n"):
        key_value = line.split(": ") 
        if len(key_value) == 2:  
            key, value = key_value  
            key = key.strip("['']")

            template = re.sub(f"{{{key}}}", value, template)
            results1[key] = value  
        else:
            
            print(f"There is a line in your file that looks like this : {line}")


while True:
    try:
        num_vars = int(
            input(
                "Enter the number of variables you want to create (enter 0 if none/skipp): "
            )
        )
        print("\n")
        break
    except ValueError:
        print("Invalid input. Please enter a Number. (0-9)")
        print("\n")

for i in range(num_vars):
    variable_name = ""
    while not is_valid_variable_name(variable_name):
        variable_name = input("Enter the name of the variable: ")
        if not is_valid_variable_name(variable_name):
            print(
                "Please enter a valid variable name. Variable names must not contain spaces or special characters, and cannot begin with a digit. Examples of acceptable variable names: 'my_variable', 'count', 'price_per_unit'."
            )
            print("\n")
    try:
        value = searchMsWord()
        results[variable_name] = value
    except IndexError:
        print(
            "Sory but the Index is not found, check for symbols attached to word and make sure to grab those also !."
        )


print(results)


def scann_variables(template_file):
    doc = docx.Document(template_file)

    variables = []
    for paragraph in doc.paragraphs:
        text = "".join(run.text for run in paragraph.runs)
        while "{{" in text and "}}" in text:
            start_index = text.index("{{")
            end_index = text.index("}}") + 2
            variable_name = text[start_index + 2 : end_index - 2]
            variables.append(variable_name)
            text = text[end_index:]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = "".join(run.text for run in paragraph.runs)
                    while "{{" in text and "}}" in text:
                        start_index = text.index("{{")
                        end_index = text.index("}}") + 2
                        variable_name = text[start_index + 2 : end_index - 2]
                        variables.append(variable_name)
                        text = text[end_index:]
    return variables


if args.output:
    if not output_filename.endswith(".txt"):
        output_filename += ".txt"
    with open(output_filename, "w") as f:
        for key, value in results.items():
            f.write(f"{key}: {value}\n")
        print("\n")
        print(f"Your Output file is successfully created !")
        print("\n")
else:
    for key, value in results.items():
        print(f"{key}: {value}")


def replace_variables(template_file, results, bold):
    variables = list(results.keys())

    doc = docx.Document(template_file)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "{{" in run.text and "}}" in run.text:
                start_index = run.text.index("{{")
                end_index = run.text.index("}}") + 2
                variable_name = run.text[start_index + 2 : end_index - 2]
                if variable_name in variables:
                    run.text = (
                        run.text[:start_index]
                        + str(results[variable_name]).strip("[']")
                        + run.text[end_index:]
                    )
                   
                    if bold:
                        run.font.bold = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if "{{" in run.text and "}}" in run.text:
                            start_index = run.text.index("{{")
                            end_index = run.text.index("}}") + 2
                            variable_name = run.text[start_index + 2 : end_index - 2]
                            if variable_name in variables:
                                run.text = (
                                    run.text[:start_index]
                                    + str(results[variable_name]).strip("[']")
                                    + run.text[end_index:]
                                )
                                
                                if bold:
                                    run.font.bold = True

    modified_template = os.path.join(
        os.getcwd(), "modified_template_{}.docx".format(os.getpid())
    )

    doc.save(modified_template)
    return modified_template


if args.template:
    if not template_filename.endswith(".docx"):
        template_filename += ".docx"
    for root, dirs, files in os.walk(search_directory):
        if template_filename in files:
            template_file_path = os.path.join(root, template_filename)
            template_file_variable = template_file_path
            break
    else:
        print(
            f"Error: {template_filename} was not found in directory on the specified {search_directory}"
        )
        print("Check the filename or check the path if it's written correctly ! ")
        print("Exiting ...")
        sys.exit()

    template_document = template_file_variable
    try:
        print(template_document, end="\n")
        document = docx.Document(template_document)
        print("\n")
    except Exception:
        print("Error: Invalid file path")
    if args.input:
        modified_template = replace_variables(template_document, results1, bold)

        print(f"Your Word Template  is successfully created!")
        print("\n")
        

    else:
        modified_template = replace_variables(template_document, results, bold)

        print(f"Your Word Template  is successfully created!")
        print("\n")

if args.scann:
    scanned_var = scann_variables(template_document)

    if args.output:
        if not output_filename.endswith(".txt"):
            output_filename += ".txt"

        with open(output_filename, "w") as f:
            for variable in scanned_var:
                f.write(f"{variable}: ['']\n")
            print("\n")
            print(f"Your Output file is successfully created !")
            print("\n")

if args.pdf:
    if args.input:
        pdf_file = f"modified_template_{os.getpid()}.pdf"
        docx2pdf.convert(modified_template, output_path=pdf_file)

        print(f"Your Word Template  is successfully created!")
        print("\n")
 
    else:
        pdf_file = f"modified_template_{os.getpid()}.pdf"
        docx2pdf.convert(modified_template, output_path=pdf_file)
        print("\n")
        print(f" Your PDF file is successfully created !")
        print("\n")


print("""¯\_( ͠❛ ͜ʖ ͠❛ )_/¯""")
